import streamlit as st
import fitz  # type: ignore # PyMuPDF
import docx # type: ignore
import os
import json
import pyttsx3 # type: ignore
import speech_recognition as sr # type: ignore
from io import BytesIO
from datetime import date, datetime
from langchain_groq import ChatGroq
import requests
from fpdf import FPDF # type: ignore
from docx import Document # type: ignore
from typing import List, Dict, Any
import base64  # Import base64 for image encoding

# ================================
# SETUP
# ================================

st.set_page_config(page_title="Elderly Support App", layout="wide")

# Initialize API Key for Groq
os.environ["GROQ_API_KEY"] = "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"
llm = ChatGroq(model="llama3-70b-8192")

# Initialize Session State
if 'family_data' not in st.session_state:
    st.session_state.family_data = []
if 'calendar_events' not in st.session_state:
    st.session_state.calendar_events = []
if 'story_text' not in st.session_state:
    st.session_state.story_text = ""
if 'is_recording' not in st.session_state:
    st.session_state.is_recording = False
if 'user_knowledgebase' not in st.session_state:
    st.session_state.user_knowledgebase = []
if 'sidebar_option' not in st.session_state:
    st.session_state.sidebar_option = "Home"  # Default to Home
if 'name' not in st.session_state:
    st.session_state.name = ""
if 'memory' not in st.session_state:
    st.session_state.memory = ""
if 'uploaded_image' not in st.session_state:
    st.session_state.uploaded_image = None
if 'uploaded_voice' not in st.session_state:
    st.session_state.uploaded_voice = None
if 'search_query' not in st.session_state:
    st.session_state.search_query = ""
if 'event_name' not in st.session_state:
    st.session_state.event_name = ""
if 'event_date' not in st.session_state:
    st.session_state.event_date = date.today()
if 'ai_response' not in st.session_state:
    st.session_state.ai_response = ""
if 'uploaded_animation_images' not in st.session_state:
    st.session_state.uploaded_animation_images = []

# Short names for pages for voice navigation
page_shortnames = {
    "home": "Home",
    "upload": "Upload Family Photos & Memories",
    "calendar": "Memory Calendar",
    "search": "Search Family Members",
    "delete": "Delete Family Member",
    "reset": "Reset Current Session Data",
    "story": "Tell Your Story",
    "ai": "Interact with AI",
    "about": "About Us"
}

# ================================
# Helper Functions
# ================================

def show_family_gallery(family_data: List[Dict[str, Any]]):
    """Displays family members with images and memories."""
    for member in family_data:
        st.subheader(member['name'])
        if isinstance(member['image'], bytes):
            st.image(member['image'], caption=f"{member['name']}", use_container_width=True)
        else:
            st.write(f"Image: {member['image']}")  # print the file name
        st.write(f"Memory: {member['memory']}")


def show_memory_backstory(family_data: List[Dict[str, Any]]):
    """Displays memories of family members."""
    st.subheader("Memories")
    if family_data:
        for member in family_data:
            st.write(f"{member['name']}: {member['memory']}")
    else:
        st.write("No family members added yet.")


def show_memory_calendar(calendar_events: List[Dict[str, str]]):
    """Displays events from the calendar."""
    st.subheader("Upcoming Events")
    if calendar_events:
        for event in calendar_events:
            st.write(f"{event['event']} on {event['date']}")
    else:
        st.write("No events added yet.")

# ================================
# VOICE COMMAND FUNCTIONS
# ================================

def recognize_speech():
    """Recognizes speech from the microphone."""
    recognizer = sr.Recognizer()
    with sr.Microphone() as source:
        recognizer.adjust_for_ambient_noise(source)
        st.info("Listening...")
        try:
            audio = recognizer.listen(source, timeout=5)  # 5-second timeout
            st.info("Recognizing...")
            text = recognizer.recognize_google(audio)
            st.success(f"You said: {text}")
            return text.lower()
        except sr.WaitTimeoutError:
            st.error("No speech detected. Please try again.")
            return ""  # Return empty string for no speech
        except sr.UnknownValueError:
            st.error("Could not understand audio")
            return ""
        except sr.RequestError as e:
            st.error(f"Could not request results from Google Speech Recognition service; {e}")
            return ""
        except Exception as e:
            st.error(f"An unexpected error occurred: {e}")
            return ""


def speak(text):
    """Speaks the given text using a text-to-speech engine."""
    engine = pyttsx3.init()
    engine.say(text)
    engine.runAndWait()


def navigate_page(page_name):
    """Navigates to the specified page."""
    # st.write(f"Navigating to: {page_name}")
    # st.write(f"Page names: {page_shortnames}")
    for key, value in page_shortnames.items():
        if page_name in value.lower() or page_name in key.lower():
            st.session_state.sidebar_option = value
            return True
    return False


def process_voice_command(command):
    """Processes the voice command and performs the corresponding action."""
    if not command:
        return  # Exit if the command is empty

    if "upload" in command and "file" in command:
        st.info("Please specify the file name you want to upload.")
        file_name = recognize_speech()
        if file_name:
            st.session_state.uploaded_file = file_name
            st.success(f"File {file_name} will be uploaded.")
        else:
            st.error("No file name provided.")

    elif "set name" in command:
        st.info("Please tell me the name.")
        name = recognize_speech()
        if name:
            st.session_state.name = name
            st.success(f"Name set to {name}")
        else:
            st.error("No name provided.")

    elif "set memory" in command:
        st.info("Please tell me the memory.")
        memory = recognize_speech()
        if memory:
            st.session_state.memory = memory
            st.success(f"Memory set to: {memory}")
        else:
            st.error("No memory provided.")

    elif "add family member" in command:
        if st.session_state.name and st.session_state.memory and st.session_state.uploaded_file:
            st.session_state.family_data.append({
                "name": st.session_state.name,
                "image": st.session_state.uploaded_file,
                "memory": st.session_state.memory,
                "voice": ""
            })
            st.success(f"{st.session_state.name} added to memory gallery")
            st.session_state.name = ""
            st.session_state.memory = ""
            st.session_state.uploaded_image = None
            st.session_state.uploaded_voice = None
        else:
            st.error("Please provide name, memory, and image before adding a family member.")

    elif "search" in command and "family" in command:
        st.info("Please tell me the name of the family member to search for.")
        search_query = recognize_speech()
        if search_query:
            st.session_state.search_query = search_query
            st.success(f"Searching for {search_query}")
        else:
            st.error("No search query provided.")

    elif "delete" in command and "family" in command:
        st.info("Please tell me the name of the family member to delete.")
        delete_name = recognize_speech()
        if delete_name:
            before = len(st.session_state.family_data)
            st.session_state.family_data = [m for m in st.session_state.family_data if
                                            m['name'].lower() != delete_name.lower()]
            if len(st.session_state.family_data) < before:
                st.success(f"Deleted {delete_name}")
            else:
                st.warning(f"{delete_name} not found")
        else:
            st.error("No name provided for deletion.")

    elif "reset session" in command:
        reset_session_data()

    elif "record story" in command:
        st.session_state.is_recording = True
        st.session_state.story_text = recognize_speech()
        st.session_state.is_recording = False

    elif "add event" in command:
        st.info("Please tell me the event name")
        event_name = recognize_speech()
        if event_name:
            st.session_state.event_name = event_name
            st.info("Please tell me the event date")
            event_date_str = recognize_speech()
            if event_date_str:
                try:
                    event_date = datetime.strptime(event_date_str, "%d-%m-%Y").date()
                    st.session_state.event_date = event_date
                    st.session_state.calendar_events.append(
                        {"event": st.session_state.event_name, "date": str(st.session_state.event_date)})
                    st.success(f"Added {st.session_state.event_name} on {st.session_state.event_date}")
                except ValueError:
                    st.error("Invalid date format.  Please use %d-%m-%Y")
            else:
                st.error("No event date provided")
        else:
            st.error("No event name provided.")

    elif "clear knowledge" in command:
        st.session_state.user_knowledgebase = []
        st.success("AI Knowledge Base Cleared")

    elif "go to" in command:
        page_name = command.replace("go to ", "").strip()
        if navigate_page(page_name):
            st.rerun()
        else:
            st.error(f"Page '{page_name}' not found.")

    elif "add date" in command:
        date_str = command.replace("add date", "").strip()
        try:
            event_date = datetime.strptime(date_str, "%Y-%m-%d").date()
            st.session_state.event_date = event_date
            st.success(f"Date set to {event_date}")
        except ValueError:
            st.error("Invalid date format. Please use %Y-%m-%d")

    elif "add event name" in command:
        event_name = command.replace("add event name", "").strip()
        st.session_state.event_name = event_name
        st.success(f"Event name set to {event_name}")

    elif "speak" in command and "response" in command:
        if 'ai_response' in st.session_state:
            speak(st.session_state.ai_response)

    else:
        st.error(f"Command '{command}' not recognized.  Please try again.")


# ================================
# RESET SESSION DATA FUNCTION
# ================================
def reset_session_data():
    """Resets all session state variables."""
    for key in list(st.session_state.keys()):
        del st.session_state[key]
    st.success("All session data has been reset.")


# ================================
# Sidebar Navigation and Voice Control Buttons
# ================================
# Create the list of page options
page_options = [
    "Home",  # Added "Home" to the page options
    "Upload Family Photos & Memories",
    "Memory Calendar",
    "Search Family Members",
    "Delete Family Member",
    "Reset Current Session Data",
    "Tell Your Story",
    "Interact with AI",
    "About Us"
]

if 'sidebar_option' not in st.session_state:
    st.session_state.sidebar_option = page_options[0]

# Create the radio button
option = st.sidebar.radio(
    "Navigate",
    page_options,
    index=page_options.index(st.session_state.sidebar_option),
)

# Add separate buttons for voice control
if st.sidebar.button("🎙️ Navigate Pages"):
    voice_command = recognize_speech()
    if voice_command:
        navigated = navigate_page(voice_command)
        if navigated:
            st.rerun()
        else:
            st.error("Page not found")

if st.sidebar.button("🎙️ Other Commands"):
    voice_command = recognize_speech()
    if voice_command:
        process_voice_command(voice_command)
        st.rerun()

# ================================
# Text to Speech
# ================================
def speak(text):
    engine = pyttsx3.init()
    engine.say(text)
    engine.runAndWait()


# ================================
# HOME PAGE
# ================================
if option == "Home":
    st.header("Welcome to Your Digital Memory Album")
    st.write("This app helps you cherish and relive your precious family memories.")

    # Option to upload images for animation
    uploaded_animation_images = st.file_uploader("Upload images for a home screen animation (optional)",
                                                 type=["jpg", "jpeg", "png"], accept_multiple_files=True)

    if uploaded_animation_images:
        st.session_state.uploaded_animation_images = [img.read() for img in uploaded_animation_images]

    # Display animated images if available
    if st.session_state.uploaded_animation_images:
        st.subheader("Your Animated Memories")
        num_images = len(st.session_state.uploaded_animation_images)
        if num_images > 0:
            # Basic fade-in/out animation using CSS
            animation_speed = 3  # seconds per image
            st.markdown(
                f"""
                <style>
                .image-container {{
                    position: relative;
                    width: 100%;
                    height: 300px; /* Adjust as needed */
                }}
                .animated-image {{
                    position: absolute;
                    top: 0;
                    left: 0;
                    width: 100%;
                    height: 100%;
                    object-fit: contain;
                    opacity: 0;
                    animation: fade {animation_speed * num_images}s infinite;
                }}
                @keyframes fade {{
                    0% {{ opacity: 0; }}
                    {(100 / num_images) * 10}% {{ opacity: 1; }}
                    {(100 / num_images) * 90}% {{ opacity: 1; }}
                    100% {{ opacity: 0; }}
                }}
                """,
                unsafe_allow_html=True
            )
            st.markdown("<div class='image-container'>", unsafe_allow_html=True)
            for i, img_bytes in enumerate(st.session_state.uploaded_animation_images):
                img_base64 = base64.b64encode(img_bytes).decode('utf-8')
                animation_delay = i * animation_speed  # Stagger the start of the fade
                st.markdown(
                    f"<img src='data:image/jpeg;base64,{img_base64}' class='animated-image' style='animation-delay: {animation_delay}s;'>",
                    unsafe_allow_html=True
                )
            st.markdown("</div>", unsafe_allow_html=True)
        else:
            st.write("Please upload some images to see the animation.")
    elif st.session_state.family_data:
        st.subheader("Your Memories")
        cols = st.columns(3)
        for i, member in enumerate(st.session_state.family_data):
            with cols[i % 3]:
                st.markdown(
                    f"<div style='border: 1px solid #e0e0e0; border-radius: 8px; padding: 8px; margin-bottom: 16px; background-color: #f9f9f9; box-shadow: 0 2px 4px rgba(0,0,0,0.1);'>\n"
                    f"    <h4 style='margin-bottom: 8px;'>{member['name']}</h4>\n",
                    unsafe_allow_html=True
                )
                if isinstance(member['image'], bytes):
                    img_base64 = base64.b64encode(member['image']).decode('utf-8')
                    st.markdown(
                        f"    <img src='data:image/jpeg;base64,{img_base64}' style='width: 100%; border-radius: 4px; margin-bottom: 8px;'>\n",
                        unsafe_allow_html=True
                    )
                else:
                    st.markdown(
                        f"    <p>Image: {member['image']}</p>\n",
                        unsafe_allow_html=True
                    )
                st.markdown(
                    f"    <p style='font-size: 0.9em; color: #555;'>{member['memory']}</p>\n"
                    f"</div>",
                    unsafe_allow_html=True
                )
    else:
        st.write("No memories added yet. Start by uploading photos and memories in the 'Upload Family Photos & Memories' section, or upload images for a home screen animation above.")

    # Subtle animation on the main header
    st.markdown(
        """
        <style>
        @keyframes subtle-pulse {
            0% { transform: scale(1); opacity: 0.9; }
            50% { transform: scale(1.02); opacity: 1; }
            100% { transform: scale(1); opacity: 0.9; }
        }
        .main-header {
            animation: subtle-pulse 3s infinite alternate;
        }
        </style>
        """,
        unsafe_allow_html=True
    )
    st.markdown("<h1 class='main-header'>Welcome</h1>", unsafe_allow_html=True)

# ================================
# UPLOAD FAMILY MEMORIES
# ================================
elif option == "Upload Family Photos & Memories":
    st.header("Upload Family Member Information")

    uploaded_image = st.file_uploader("Upload a photo", type=["jpg", "jpeg", "png"])
    name = st.text_input("Name", value=st.session_state.name)
    memory = st.text_area("Enter a memory", value=st.session_state.memory)
    uploaded_voice = st.file_uploader("Upload a voice recording", type=["mp3", "wav", "ogg"])

    if st.button("Add Family Member"):
        if uploaded_image and name and memory and uploaded_voice:
            voice_bytes = uploaded_voice.read()
            st.session_state.family_data.append({
                "name": name,
                "image": uploaded_image.read(),
                "memory": memory,
                "voice": voice_bytes
            })
            st.success(f"{name} added to memory gallery")
            st.session_state.name = ""
            st.session_state.memory = ""
            st.session_state.uploaded_image = None
            st.session_state.uploaded_voice = None
        else:
            st.error("Please provide all details")

    if st.session_state.family_data:
        st.subheader("Family Gallery")
        show_family_gallery(st.session_state.family_data)

    if st.session_state.family_data:
        show_memory_backstory(st.session_state.family_data)


# ================================
# MEMORY CALENDAR
# ================================
elif option == "Memory Calendar":
    st.header("Memory Calendar")
    event = st.text_input("Event Name", value=st.session_state.event_name)
    event_date = st.date_input("Event Date", value=st.session_state.event_date)
    if st.button("Add Event"):
        st.session_state.calendar_events.append({"event": event, "date": str(event_date)})
        st.success(f"Added {event} on {event_date}")
        st.session_state.event_name = ""
        st.session_state.event_date = date.today()

    if st.session_state.calendar_events:
        show_memory_calendar(st.session_state.calendar_events)


# ================================
# SEARCH FAMILY
# ================================
elif option == "Search Family Members":
    st.header("Search Family")
    search = st.text_input("Enter name to search", value=st.session_state.search_query)
    if search:
        results = [m for m in st.session_state.family_data if search.lower() in m['name'].lower()]
        if results:
            show_family_gallery(results)
            show_memory_backstory(results)
        else:
            st.warning("No match found")
    else:
        st.warning("Please enter a name to search.")


# ================================
# DELETE FAMILY
# ================================
elif option == "Delete Family Member":
    st.header("Delete Family Member")
    name_to_delete = st.text_input("Enter name to delete")
    if st.button("Delete"):
        before = len(st.session_state.family_data)
        st.session_state.family_data = [m for m in st.session_state.family_data if
                                        m['name'].lower() != name_to_delete.lower()]
        if len(st.session_state.family_data) < before:
            st.success(f"Deleted {name_to_delete}")
        else:
            st.warning(f"{name_to_delete} not found")



# ================================
# RESET SESSION
# ================================
elif option == "Reset Current Session Data":
    # Clear all session state variables, except for the persistent data
    st.header("Reset Current Session Data")
    st.write("Click the button below to reset all session data (except the family data or other important data).")

    # Button to reset session data without removing important persistent data
    if st.button("Reset All Data"):
        # List the session state keys you want to retain (important persistent data)
        persistent_data_keys = ['family_data']  # Add other keys here that should not be reset

        # Delete all session state variables except persistent ones
        for key in list(st.session_state.keys()):
            if key not in persistent_data_keys:
                del st.session_state[key]

        st.success("Current session data has been reset. Persistent data is retained.")


# ================================
# Function to recognize speech
# ================================
def recognize_speech():
    recognizer = sr.Recognizer()
    with sr.Microphone() as source:
        recognizer.adjust_for_ambient_noise(source)
        st.info("Listening... Speak now!")
        try:
            audio = recognizer.listen(source)
            st.info("Recognizing...")
            text = recognizer.recognize_google(audio)
            st.success("Transcription Complete!")
            return text
        except sr.UnknownValueError:
            st.error("Could not understand audio")
            return None
        except sr.RequestError:
            st.error("Could not request results from Google Speech Recognition service")
            return None



# ================================
# TELL YOUR STORY
# ================================
if option == "Tell Your Story":
    st.header("Record Your Story")
    st.write("Click below and start speaking about your childhood.")

    if 'is_recording' not in st.session_state:
        st.session_state.is_recording = False

    if st.button("Record Story") and not st.session_state.is_recording:
        st.session_state.is_recording = True
        st.session_state.story_text = ""
        story = recognize_speech()
        if story:
            st.session_state.story_text = story
            st.success("Story Recorded")
            st.text_area("Transcribed Story", value=story, height=200)
            with open("story_memory.txt", "a", encoding="utf-8") as f:
                f.write(story + "\n\n")
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.set_font("Arial", size=12)
            for line in story.split('\n'):
                pdf.multi_cell(0, 10, line)
            pdf.output("story_memory.pdf")
            doc = Document()
            doc.add_heading('Transcribed Story', 0)
            for para in story.split('\n'):
                doc.add_paragraph(para)
            doc.save("story_memory.docx")
        st.session_state.is_recording = False

    elif st.session_state.is_recording:
        st.info("Recording in progress... Please speak now.")

    if st.session_state.story_text:
        if os.path.exists("story_memory.txt"):
            with open("story_memory.txt", "rb") as f_txt:
                st.download_button("📥 Download as TXT", f_txt, file_name="story_memory.txt")

        if os.path.exists("story_memory.pdf"):
            with open("story_memory.pdf", "rb") as f_pdf:
                st.download_button("📥 Download as PDF", f_pdf, file_name="story_memory.pdf")

        if os.path.exists("story_memory.docx"):
            with open("story_memory.docx", "rb") as f_docx:
                st.download_button("📥 Download as DOCX", f_docx, file_name="story_memory.docx")



# ================================
# INTERACT WITH AI
# ================================
if option == "Interact with AI":
    st.header("🤖 Chat with Your Personalized AI")

    if 'user_knowledgebase' not in st.session_state:
        st.session_state.user_knowledgebase = []

    uploaded_file = st.file_uploader("Upload a document to create a personalized AI", type=["txt", "pdf", "docx"])

    if uploaded_file is not None:
        if uploaded_file.type == "application/pdf":
            doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
            file_contents = "\n".join([page.get_text() for page in doc])
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            docx_file = docx.Document(uploaded_file)
            file_contents = "\n".join([para.text for para in docx_file.paragraphs])
        else:
            file_contents = uploaded_file.read().decode("utf-8")

        st.session_state.user_knowledgebase.append(file_contents)
        st.success("✅ File uploaded and context stored for AI.")

    user_input = st.text_input("💬 Ask something or tell your story:")
    if st.button("🎙️"):
        voice_input = recognize_speech()
        if voice_input:
            user_input = voice_input # updates the user_input
            st.write(f"You said: {user_input}")

    if user_input:
        full_context = "\n\n".join(st.session_state.user_knowledgebase)
        prompt = f"You are a helpful assistant. Here's the user's document context:\n\n{full_context}\n\nNow answer the following:\n{user_input}"

        response = requests.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {os.environ['GROQ_API_KEY']}",
                "Content-Type": "application/json"
            },
            json={
                "model": "llama3-70b-8192",
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.7
            }
        )

        if response.status_code == 200:
            ai_response = response.json()["choices"][0]["message"]["content"]
            st.session_state.ai_response = ai_response
            st.write("🧠 AI says:", ai_response)

            if st.button("🔊 Speak AI Response"):
                speak(ai_response)
        else:
            st.error(f"❌ Error: {response.text}")

    if 'story_text' in st.session_state and st.session_state.story_text: # change here
        st.text_area("📝 Your Earlier Recorded Story", value=st.session_state.story_text, height=200)

        if os.path.exists("story_memory.txt"):
            with open("story_memory.txt", "rb") as f:
                st.download_button("📥 Download Your Stories", f, file_name="story_memory.txt")

        if os.path.exists("story_memory.pdf"):
            with open("story_memory.pdf", "rb") as f_pdf:
                st.download_button("📥 Download as PDF", f_pdf, file_name="story_memory.pdf")

        if os.path.exists("story_memory.docx"):
            with open("story_memory.docx", "rb") as f_docx:
                st.download_button("📥 Download as DOCX", f_docx, file_name="story_memory.docx")

# ================================
# ABOUT
# ================================
elif option == "About Us":
    st.header("About")
    st.write("This app helps elderly users connect to family memories using voice, photos, and AI conversation.")
